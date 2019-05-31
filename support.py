import numpy as np
import PyPDF2
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx import Document
import pymysql
from bs4 import BeautifulSoup  
import requests, io, re
import string
from nltk.tokenize import word_tokenize, TreebankWordTokenizer
from nltk.corpus import stopwords, wordnet
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from nltk.stem import WordNetLemmatizer
from sklearn.metrics.pairwise import cosine_similarity
from gensim.models import KeyedVectors, Word2Vec
from scipy import spatial

def readFile(doc):
    '''
    Read a text file and return to a elements of a list
    '''
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
        
    return '\n'.join(fullText)    

def onlRefs(rels):
    '''
    Extract all the hyperlinks (references) in the document
    '''
    links = []
    for rel in rels:
        if rels[rel].reltype == RT.HYPERLINK:
            links.append(rels[rel]._target)
            
    return links

def readContent(link):    
    '''
    Depend on the website (pdf or regular html)
    Open the file and scrape the data of 01 site
    Libraries used: 
    bs4 -> BeautifulSoup 
    requests, io, re
    '''
    
    string = []
   
    # if the link is a pdf
    if (r'.pdf' in link.split('/')[-1]):
        title = link.split('/')[-1]
        response = requests.get(link)
        raw_data = response.content
        pdf_content = io.BytesIO(raw_data)
        pdf_reader = PyPDF2.PdfFileReader(pdf_content)
        for page in range(pdf_reader.numPages):
            string.append(pdf_reader.getPage(page).extractText())
        return link, title, (' '.join(string))
    
    # if not
    else:
        def scrape_data():
            '''
            Return title + content of a webpage 
            '''
            page = requests.get(link)
            title = BeautifulSoup(page.text, 'html.parser').head.title.contents
            text = BeautifulSoup(page.text, 'html.parser').find_all('p')
            for p in text:
                string.append(p.get_text())
            return link, title, (' '.join(string).replace(u'\xa0', ' ').replace(u'\n', ' '))
        
        try:
            return scrape_data()
        
        #some links need authentication
        except:
            
            headers = {'User-Agent':'Mozilla/5.0'}             
            #class AppURLopener(urllib.request.FancyURLopener):
                #version = "Mozilla/5.0"
            #opener = AppURLopener()
            return scrape_data()

def working_with_mySQL(readContent_result):
    '''
    Store (link, title, content) scraped from 01 website to local db
    Return content
    Library used: pymysql
    '''
    #Open database
    db = pymysql.connect(host = '127.0.0.1',
                          user = 'root',
                           db = 'references')
    
    cursor = db.cursor()   
    
    #check distinct reference to add to the database
    sql = "INSERT INTO onlref (link, title, content) VALUES (%s, %s, %s)"

    try:
        cursor.execute(sql, readContent_result)
    except:
        pass

    #fetch data
    link, title, content = readContent_result
    cursor.execute('SELECT content FROM onlref WHERE onlref.link = %s', link)
    data = cursor.fetchone()

    db.commit()
    db.close()
        
    return str(data[0])

def getToken(text):
    '''
    Tokenise + Omit punctuation
    Libraries: 
        nltk.tokenize -> word_tokenize, 
        nltk.corpus -> stopwords,
        string
    '''
    translator = str.maketrans('', '', string.punctuation)
    stopWords = set(stopwords.words('english'))
    tokens = word_tokenize(text)
    tokens = [token.lower() for token in tokens]
    filtered = [w for w in tokens if not w in stopWords]
    
    return ' '.join(filtered).translate(translator).split()

def lemmatize_append(set_tokens):
    lmt = WordNetLemmatizer()
    lemmatize_set = []
    for wd in set_tokens:
        lemmatize_set.append(lmt.lemmatize(wd))
    return set(lemmatize_set)

class Methods(object):
    def __init__(self, doc):
        self.doc = doc
        
    def jaccard_similarity(self):
        '''
        Checking similarity using Jaccard Similarity
        '''
        
        document = Document(self.doc)
        rels = document.part.rels
        tokens_a = set(getToken(readFile(document)))
        features = []
        percentage = []
        
        def jaccard(a, b, j):
            return float(len(j)/(len(a) + len(b) - len(j)))

        features.append(lemmatize_append(tokens_a))

        for link in onlRefs(rels):
            features.append(lemmatize_append(getToken(working_with_mySQL(readContent(link)))))
            
        while len(features) > 2:
            i = len(features) - 1
            features[1] = features[1].union(features[i])
            features.pop(i)
            
        jac = features[0].intersection(features[1])

        return jaccard(features[0], features[1], jac)*100
    
    def cosine_sim(self):
        '''
        Perform Cosine Similarity
        '''

        tokenizer = TreebankWordTokenizer()
        vect = CountVectorizer()
        vect.set_params(tokenizer=tokenizer.tokenize, stop_words='english')
        document = Document(self.doc)
        rels = document.part.rels
        
        corpus = []
        corpus.append(readFile(document))

        for link in onlRefs(rels):
            corpus.append(working_with_mySQL(readContent(link)))

        tfidf = vect.fit_transform(corpus)
        return (1 - cosine_similarity(tfidf)[0][1])*100
    
    def word_to_vec(self):
        '''
        Comparing the contextual similarity between documents.
        Converting tokens to numeric vector using Google pretrained document.
        Perform cosine similarity based on that.
        '''
        
        #take only first 100k most frequent tokens
        model = KeyedVectors.load_word2vec_format('GoogleNews-vectors-negative300.bin', binary=True, limit = 100000)
        
        document = Document(self.doc)
        rels = document.part.rels
        result = []
        
        def convert(corpora):
            return np.mean([model[wd] for wd in getToken(corpora) if wd in model], axis=0)
            
        base_corpora = convert(readFile(document))
        for link in onlRefs(rels):
            cal = 1 - spatial.distance.cosine(
                base_corpora, 
                convert(working_with_mySQL(readContent(link))))
            
            result.append(cal)
            
        return np.mean(result)*100